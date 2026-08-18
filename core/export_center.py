import base64
import io
import json
import re
import zipfile
from datetime import datetime
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from core.runtime import DB_PATH, read_df
from core.portfolio_center import portfolio_ranking, portfolio_summary, money_short

EXPORT_VERSION = "V0.6.1"

PROJECT_TABLES = [
    "projects", "work_items", "contracts", "materials", "labour_logs",
    "costs", "daily_logs", "changes", "acceptances", "payments",
    "warranty_maintenance", "attachments"
]
ALL_TABLES = [
    "projects", "partners", "contracts", "work_items", "materials",
    "labour_logs", "costs", "daily_logs", "changes", "acceptances",
    "payments", "warranty_maintenance", "attachments"
]


WORD_LABELS = {
    "work_id": "Mã việc", "work_name": "Công việc", "planned_finish": "Hạn KH",
    "planned_progress": "Tiến độ KH", "actual_progress": "Tiến độ TT", "status": "Trạng thái",
    "assignee": "Phụ trách", "cost_group": "Nhóm chi phí", "description": "Nội dung",
    "planned_budget": "Ngân sách KH", "actual_cost": "Chi thực tế", "forecast_final": "Dự báo cuối kỳ",
    "change_id": "Mã PS", "found_date": "Ngày phát hiện", "proposed_value": "Giá trị đề xuất",
    "approval_status": "Trạng thái duyệt", "approved_value": "Giá trị duyệt",
    "acceptance_id": "Mã NT", "acceptance_date": "Ngày nghiệm thu", "result": "Kết quả",
    "defects": "Lỗi/Tồn tại", "correction_due": "Hạn khắc phục", "accepted_value": "Giá trị NT",
    "payment_id": "Mã TT", "tranche": "Đợt thanh toán", "paid_value": "Đã thanh toán",
    "payment_date": "Ngày thanh toán", "wm_id": "Mã BH/BT", "asset_item": "Hạng mục/Tài sản",
    "wm_type": "Loại", "due_date": "Đến hạn", "cost": "Chi phí",
    "attachment_id": "Mã hồ sơ", "upload_date": "Ngày tải", "category": "Loại hồ sơ",
    "filename": "Tên tệp", "uploaded_by": "Người tải",
    "project_id": "Mã dự án", "project_name": "Tên dự án", "project_type": "Loại dự án",
    "owner": "Chủ đầu tư/Đơn vị", "location": "Địa điểm", "budget": "Ngân sách",
}
MONEY_COLUMNS = {
    "planned_budget", "actual_cost", "forecast_final", "proposed_value", "approved_value",
    "accepted_value", "paid_value", "cost", "budget", "contract_value", "advance_value",
    "requested_value", "committed", "unit_price"
}
DATE_COLUMNS = {
    "planned_finish", "found_date", "acceptance_date", "correction_due", "payment_date",
    "due_date", "upload_date", "planned_start", "request_date", "signed_date", "start_date",
    "finish_date", "work_date", "log_date"
}
PERCENT_COLUMNS = {"planned_progress", "actual_progress"}

TABLE_LABELS = {
    "projects": "Dự án",
    "partners": "Đối tác",
    "contracts": "Hợp đồng",
    "work_items": "Công việc",
    "materials": "Vật tư",
    "labour_logs": "Nhân công",
    "costs": "Chi phí",
    "daily_logs": "Nhật ký",
    "changes": "Phát sinh",
    "acceptances": "Nghiệm thu",
    "payments": "Thanh toán",
    "warranty_maintenance": "Bảo hành_Bảo trì",
    "attachments": "Hồ sơ đính kèm",
}


def _safe_name(name):
    text = str(name or "file").strip()
    text = re.sub(r'[\\/:*?"<>|]+', '_', text)
    text = re.sub(r'\s+', '_', text)
    return text[:120] or "file"


def _now_stamp():
    return datetime.now().strftime("%Y%m%d_%H%M")


def _table_exists(table):
    df = read_df(
        "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
        (table,)
    )
    return not df.empty


def _project_partner_ids(pid):
    if not _table_exists("contracts"):
        return []
    df = read_df(
        "SELECT DISTINCT partner_id FROM contracts WHERE project_id=? AND partner_id IS NOT NULL AND partner_id<>''",
        (pid,)
    )
    return df["partner_id"].tolist() if not df.empty else []


def _clean_export_df(table, df):
    out = df.copy()
    if table == "attachments" and "content_b64" in out.columns:
        out = out.drop(columns=["content_b64"])
    return out


def project_tables(pid):
    data = {}
    for table in PROJECT_TABLES:
        if not _table_exists(table):
            data[table] = pd.DataFrame()
            continue
        data[table] = _clean_export_df(
            table,
            read_df(f"SELECT * FROM {table} WHERE project_id=?", (pid,))
        )

    if _table_exists("partners"):
        partner_ids = _project_partner_ids(pid)
        if partner_ids:
            marks = ",".join(["?"] * len(partner_ids))
            data["partners"] = read_df(
                f"SELECT * FROM partners WHERE partner_id IN ({marks})", tuple(partner_ids)
            )
        else:
            data["partners"] = pd.DataFrame()
    return data


def portfolio_tables():
    data = {}
    for table in ALL_TABLES:
        if not _table_exists(table):
            data[table] = pd.DataFrame()
            continue
        data[table] = _clean_export_df(table, read_df(f"SELECT * FROM {table}"))
    return data


def _attachment_rows(pid=None):
    if not _table_exists("attachments"):
        return pd.DataFrame()
    if pid:
        return read_df("SELECT * FROM attachments WHERE project_id=?", (pid,))
    return read_df("SELECT * FROM attachments")


def _attachment_files(pid=None):
    df = _attachment_rows(pid)
    files = []
    if df.empty:
        return files
    used = set()
    for _, r in df.iterrows():
        raw = r.get("content_b64", "")
        if not isinstance(raw, str) or not raw.strip():
            continue
        try:
            content = base64.b64decode(raw)
        except Exception:
            continue
        aid = _safe_name(r.get("attachment_id", "ATT"))
        fname = _safe_name(r.get("filename", "file.bin"))
        name = f"{aid}_{fname}"
        n = 2
        original = name
        while name in used:
            stem = Path(original).stem
            suf = Path(original).suffix
            name = f"{stem}_{n}{suf}"
            n += 1
        used.add(name)
        files.append((name, content))
    return files


def _doc_setup(doc, landscape=False):
    section = doc.sections[0]
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(1.8 if landscape else 2.2)
    section.right_margin = Cm(1.8 if landscape else 2.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for style_name in ["Title", "Heading 1", "Heading 2"]:
        if style_name in styles:
            styles[style_name].font.name = "Times New Roman"
    styles["Title"].font.size = Pt(18)
    styles["Heading 1"].font.size = Pt(15)
    styles["Heading 2"].font.size = Pt(13)


def _add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(11)


def _add_kv_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text = str(k)
        cells[1].text = str(v)
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    return table


def _repeat_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _add_df_table(doc, df, columns=None, max_rows=25):
    if df is None or df.empty:
        doc.add_paragraph("Không có dữ liệu.")
        return
    work = df.copy()
    if columns:
        cols = [c for c in columns if c in work.columns]
        work = work[cols]

    for c in list(work.columns):
        if c in MONEY_COLUMNS:
            work[c] = work[c].apply(money_short)
        elif c in DATE_COLUMNS:
            work[c] = work[c].apply(
                lambda x: pd.to_datetime(x, errors="coerce").strftime("%d/%m/%Y")
                if pd.notna(pd.to_datetime(x, errors="coerce")) else ""
            )
        elif c in PERCENT_COLUMNS:
            work[c] = work[c].apply(lambda x: f"{float(x or 0):.0f}%")

    work = work.rename(columns=WORD_LABELS)
    work = work.head(max_rows).fillna("")
    table = doc.add_table(rows=1, cols=len(work.columns))
    table.style = "Table Grid"
    for i, c in enumerate(work.columns):
        table.rows[0].cells[i].text = str(c)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    _repeat_header_row(table.rows[0])
    for _, row in work.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            cells[i].text = str(value)


def project_word_bytes(pid):
    tables = project_tables(pid)
    proj = tables.get("projects", pd.DataFrame())
    if proj.empty:
        raise ValueError("Không tìm thấy dự án.")
    p = proj.iloc[0]
    ranking = portfolio_ranking()
    risk = ranking[ranking["Mã dự án"] == pid]
    r = risk.iloc[0] if not risk.empty else None

    doc = Document()
    _doc_setup(doc)
    _add_title(
        doc,
        "AI-KOS DeliveryOS – BÁO CÁO DỰ ÁN",
        f"Mã dự án: {pid} | Xuất từ Streamlit | {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )

    doc.add_heading("1. Thông tin dự án", level=1)
    _add_kv_table(doc, [
        ("Tên dự án", p.get("project_name", "")),
        ("Loại dự án", p.get("project_type", "")),
        ("Chủ đầu tư/Đơn vị", p.get("owner", "")),
        ("Địa điểm", p.get("location", "")),
        ("Trạng thái", p.get("status", "")),
        ("Ngân sách", money_short(p.get("budget", 0))),
        ("Tiến độ kế hoạch", f"{float(p.get('planned_progress',0) or 0):.0f}%"),
        ("Tiến độ thực tế", f"{float(p.get('actual_progress',0) or 0):.0f}%"),
    ])

    doc.add_heading("2. Chỉ báo điều hành", level=1)
    if r is not None:
        _add_kv_table(doc, [
            ("Điểm rủi ro", f"{int(r['Điểm rủi ro'])}/100 – {r['Mức rủi ro']}"),
            ("Dự báo cuối kỳ", money_short(r["Dự báo"])),
            ("Chênh dự báo", money_short(r["Chênh dự báo"])),
            ("Chênh tiến độ", f"{r['Chênh tiến độ']:.0f} điểm %"),
            ("Việc quá hạn", int(r["Việc quá hạn"])),
            ("Phát sinh chờ duyệt", int(r["PS chờ duyệt"])),
            ("Nghiệm thu còn lỗi", int(r["NT còn lỗi"])),
            ("Còn phải thanh toán", money_short(r["Còn phải trả"])),
        ])

    sections = [
        ("3. Công việc & tiến độ", "work_items", ["work_id","work_name","planned_finish","planned_progress","actual_progress","status","assignee"]),
        ("4. Chi phí", "costs", ["cost_group","description","planned_budget","actual_cost","forecast_final"]),
        ("5. Phát sinh", "changes", ["change_id","found_date","description","proposed_value","approval_status","approved_value"]),
        ("6. Nghiệm thu", "acceptances", ["acceptance_id","acceptance_date","result","defects","correction_due","accepted_value"]),
        ("7. Thanh toán", "payments", ["payment_id","tranche","approved_value","paid_value","status","payment_date"]),
        ("8. Bảo hành & bảo trì", "warranty_maintenance", ["wm_id","asset_item","wm_type","due_date","status","cost"]),
    ]
    for title, key, cols in sections:
        doc.add_heading(title, level=1)
        _add_df_table(doc, tables.get(key), cols, max_rows=30)

    doc.add_heading("9. Hồ sơ đính kèm", level=1)
    att = tables.get("attachments", pd.DataFrame())
    _add_df_table(doc, att, ["attachment_id","upload_date","category","filename","description","uploaded_by"], max_rows=50)

    doc.add_paragraph(
        "Ghi chú: Báo cáo được tạo tự động từ dữ liệu DeliveryOS tại thời điểm xuất. "
        "Hệ thống chỉ tổng hợp dữ liệu, không thay thế hồ sơ gốc hoặc quyết định của người có thẩm quyền."
    )
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def portfolio_word_bytes():
    ranking = portfolio_ranking()
    summary = portfolio_summary(ranking)
    doc = Document()
    _doc_setup(doc, landscape=True)
    _add_title(
        doc,
        "AI-KOS DeliveryOS – BÁO CÁO DANH MỤC DỰ ÁN",
        f"Xuất từ Streamlit | {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    )
    doc.add_heading("1. Tổng hợp danh mục", level=1)
    _add_kv_table(doc, [
        ("Số dự án", summary["projects"]),
        ("Tổng ngân sách", money_short(summary["budget"])),
        ("Tổng dự báo cuối kỳ", money_short(summary["forecast"])),
        ("Dự báo vượt ngân sách cộng dồn", money_short(summary["overrun"])),
        ("Còn phải thanh toán", money_short(summary["outstanding"])),
        ("Dự án rủi ro rất cao", summary["red"]),
        ("Dự án rủi ro cao", summary["orange"]),
        ("Công việc quá hạn", summary["overdue_work"]),
        ("Phát sinh chờ duyệt", summary["pending_changes"]),
        ("Nghiệm thu còn lỗi", summary["defects"]),
    ])
    doc.add_heading("2. Xếp hạng rủi ro & tiến độ", level=1)
    if not ranking.empty:
        show = ranking.copy()
        show.insert(0, "Xếp hạng", range(1, len(show)+1))
        _add_df_table(
            doc, show,
            ["Xếp hạng","Điểm rủi ro","Mức rủi ro","Mã dự án","Dự án","KH (%)","TT (%)","Chênh tiến độ"],
            max_rows=100
        )

        doc.add_heading("3. Tài chính & công việc cần chú ý", level=1)
        show2 = ranking.copy()
        show2.insert(0, "Xếp hạng", range(1, len(show2)+1))
        show2["Ngân sách"] = show2["Ngân sách"].apply(money_short)
        show2["Dự báo"] = show2["Dự báo"].apply(money_short)
        show2["Chênh dự báo"] = show2["Chênh dự báo"].apply(money_short)
        show2["Còn phải trả"] = show2["Còn phải trả"].apply(money_short)
        _add_df_table(
            doc, show2,
            ["Xếp hạng","Mã dự án","Ngân sách","Dự báo","Chênh dự báo","Việc quá hạn","PS chờ duyệt","NT còn lỗi","Còn phải trả"],
            max_rows=100
        )
    doc.add_heading("4. Nguyên tắc sử dụng", level=1)
    doc.add_paragraph(
        "Điểm rủi ro là chỉ báo quản trị để xếp thứ tự chú ý. Báo cáo không tự sửa dữ liệu nguồn, "
        "không tự phê duyệt phát sinh, nghiệm thu hoặc thanh toán."
    )
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _autosize_excel(writer):
    from openpyxl.styles import Font, Alignment
    for ws in writer.book.worksheets:
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for col_cells in ws.columns:
            letter = col_cells[0].column_letter
            max_len = 0
            for cell in col_cells[:200]:
                value = "" if cell.value is None else str(cell.value)
                max_len = max(max_len, len(value))
            ws.column_dimensions[letter].width = min(max(max_len + 2, 10), 36)


def _write_excel(data, overview_df=None):
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="openpyxl") as writer:
        if overview_df is not None:
            overview_df.to_excel(writer, sheet_name="Tong_quan", index=False)
        for table, df in data.items():
            sheet = TABLE_LABELS.get(table, table)[:31]
            df.to_excel(writer, sheet_name=sheet, index=False)
        _autosize_excel(writer)
    out.seek(0)
    return out.getvalue()


def project_excel_bytes(pid):
    data = project_tables(pid)
    ranking = portfolio_ranking()
    overview = ranking[ranking["Mã dự án"] == pid].copy()
    return _write_excel(data, overview)


def portfolio_excel_bytes():
    data = portfolio_tables()
    ranking = portfolio_ranking()
    return _write_excel(data, ranking)


def _records_json(data):
    payload = {}
    for table, df in data.items():
        clean = df.where(pd.notna(df), None)
        payload[table] = clean.to_dict(orient="records")
    return payload


def project_json_bytes(pid):
    data = project_tables(pid)
    payload = {
        "export_version": EXPORT_VERSION,
        "scope": "project",
        "project_id": pid,
        "exported_at": datetime.now().isoformat(timespec="seconds"),
        "tables": _records_json(data),
    }
    return json.dumps(payload, ensure_ascii=False, indent=2, default=str).encode("utf-8")


def portfolio_json_bytes():
    data = portfolio_tables()
    payload = {
        "export_version": EXPORT_VERSION,
        "scope": "portfolio",
        "exported_at": datetime.now().isoformat(timespec="seconds"),
        "ranking": portfolio_ranking().where(pd.notna(portfolio_ranking()), None).to_dict(orient="records"),
        "tables": _records_json(data),
    }
    return json.dumps(payload, ensure_ascii=False, indent=2, default=str).encode("utf-8")


def _csv_items(data):
    items = []
    for table, df in data.items():
        name = f"{table}.csv"
        content = df.to_csv(index=False).encode("utf-8-sig")
        items.append((name, content))
    return items


def raw_csv_zip_bytes(pid=None):
    data = project_tables(pid) if pid else portfolio_tables()
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        for name, content in _csv_items(data):
            z.writestr(name, content)
    return out.getvalue()


def attachments_zip_bytes(pid=None):
    files = _attachment_files(pid)
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        if files:
            for name, content in files:
                z.writestr(name, content)
        else:
            z.writestr("README.txt", "Không có hồ sơ đính kèm được lưu trong CSDL tại thời điểm xuất.")
    return out.getvalue()


def _manifest(scope, pid, data, attachment_count, include_db):
    return {
        "product": "AI-KOS DeliveryOS",
        "export_version": EXPORT_VERSION,
        "scope": scope,
        "project_id": pid,
        "exported_at": datetime.now().isoformat(timespec="seconds"),
        "table_counts": {k: int(len(v)) for k, v in data.items()},
        "attachment_files": int(attachment_count),
        "database_backup_included": bool(include_db),
        "formats": ["DOCX", "XLSX", "CSV", "JSON", "ZIP"] + (["SQLite DB"] if include_db else []),
    }


def project_full_zip_bytes(pid):
    data = project_tables(pid)
    attachments = _attachment_files(pid)
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"01_{pid}_Bao_cao.docx", project_word_bytes(pid))
        z.writestr(f"02_{pid}_Du_lieu.xlsx", project_excel_bytes(pid))
        z.writestr(f"03_{pid}_Du_lieu.json", project_json_bytes(pid))
        for name, content in _csv_items(data):
            z.writestr(f"04_Du_lieu_tho/{name}", content)
        if attachments:
            for name, content in attachments:
                z.writestr(f"05_Ho_so_dinh_kem/{name}", content)
        else:
            z.writestr("05_Ho_so_dinh_kem/README.txt", "Không có hồ sơ đính kèm.")
        manifest = _manifest("project", pid, data, len(attachments), False)
        z.writestr("00_MANIFEST.json", json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8"))
    return out.getvalue()


def portfolio_full_zip_bytes():
    data = portfolio_tables()
    attachments = _attachment_files(None)
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("01_Bao_cao_danh_muc.docx", portfolio_word_bytes())
        z.writestr("02_Du_lieu_toan_bo.xlsx", portfolio_excel_bytes())
        z.writestr("03_Du_lieu_toan_bo.json", portfolio_json_bytes())
        for name, content in _csv_items(data):
            z.writestr(f"04_Du_lieu_tho/{name}", content)
        if attachments:
            for name, content in attachments:
                z.writestr(f"05_Ho_so_dinh_kem/{name}", content)
        else:
            z.writestr("05_Ho_so_dinh_kem/README.txt", "Không có hồ sơ đính kèm.")
        if DB_PATH.exists():
            z.writestr("06_CSDL_Sao_luu/deliveryos.db", DB_PATH.read_bytes())
        manifest = _manifest("portfolio", None, data, len(attachments), DB_PATH.exists())
        z.writestr("00_MANIFEST.json", json.dumps(manifest, ensure_ascii=False, indent=2).encode("utf-8"))
    return out.getvalue()
